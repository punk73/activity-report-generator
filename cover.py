import os
import re
from datetime import datetime
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

# Fungsi untuk mendapatkan tanggal dan tempat dari nama folder
def get_date_and_place_from_folder(folder_name):
    try:
        date_str, place_info = folder_name.split(' ', 1)
        date_obj = datetime.strptime(date_str, "%Y-%m-%d")

        days = ['Senin', 'Selasa', 'Rabu', 'Kamis', 'Jumat', 'Sabtu', 'Minggu']
        months = ['empty', 'Januari', 'Februari', 'Maret', 'April', 'Mei', 'Juni', 'Juli', 'Agustus', 'September', 'Oktober', 'November', 'Desember']

        day_name = days[date_obj.weekday()]
        day = date_obj.day
        month = months[date_obj.month]
        year = date_obj.year

        formatted_date = f"{day_name}, {day} {month} {year}"

        # Memisahkan tempat dari informasi folder
        place_matches = re.split(r"\b[Dd]i\b", place_info)

        if len(place_matches) > 1:
            place = place_matches[-1].strip()

            # Periksa apakah tempat mengandung angka atau jika hanya angka saja
            if any(char.isdigit() for char in place):
                place = "Zoom Meeting"
        else:
            place = "Zoom Meeting"

        return formatted_date, place, month
    except ValueError:
        print(f"Nama folder '{folder_name}' tidak sesuai format yang diharapkan (YYYY-MM-DD Tempat).")
        return None, None, None

# Fungsi untuk mengatur border pada tabel
def set_cell_border(cell, **kwargs):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()

    for border_name, border_value in kwargs.items():
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), border_value.get('val', 'single'))
        border.set(qn('w:sz'), str(border_value.get('sz', 8)))
        border.set(qn('w:space'), str(border_value.get('space', 0)))
        border.set(qn('w:color'), border_value.get('color', '000000'))
        tcPr.append(border)

# Fungsi untuk mengisi tabel yang sudah ada
def fill_existing_table(doc, date_texts, place_texts):
    tables = doc.tables
    if len(tables) >= 2:
        target_table = tables[1]

        for row in target_table.rows[2:]:
            target_table._element.remove(row._element)

        for index, (place, date) in enumerate(zip(place_texts, date_texts), start=1):
            row_cells = target_table.add_row().cells
            row_cells[0].text = f"{index}. {place}"  # Menambahkan nomor di depan nama tempat
            row_cells[1].text = date

            for cell in row_cells:
                set_cell_border(cell,
                    top={"val": "single", "sz": 8, "color": "000000"},
                    left={"val": "single", "sz": 8, "color": "000000"},
                    bottom={"val": "single", "sz": 8, "color": "000000"},
                    right={"val": "single", "sz": 8, "color": "000000"}
                )

                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(11)

                if cell == row_cells[1]:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = 1

        for i in range(len(target_table.columns)):
            max_width = 0
            for row in target_table.rows:
                cell_text = row.cells[i].text
                max_width = max(max_width, len(cell_text))

            target_table.columns[i].width = Pt(max_width * 7)

# Fungsi untuk mengganti placeholder dengan jumlah folder yang ditemukan
def replace_placeholder_in_paragraphs(doc, folder_count):
    placeholder = "{{jumlah_folder}}"
    for para in doc.paragraphs:
        for run in para.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, str(folder_count))
                run.font.bold = True
                run.font.size = Pt(12)
                run.font.name = 'Arial'
                print(f"Placeholder '{placeholder}' berhasil diganti dengan {folder_count}.")

# Fungsi untuk mengganti nama bulan dengan font 16 di halaman pertama, dan font 12 di halaman selanjutnya
def replace_month_names(doc, month):
    found_month_placeholder = False  # Untuk melacak apakah placeholder bulan ditemukan
    is_first_page = True  # Flag untuk memeriksa apakah masih di halaman pertama
    element_count = 0  # Menghitung elemen yang diproses di halaman pertama
    
    # Mengganti bulan di paragraf
    for para in doc.paragraphs:
        for run in para.runs:
            if '{{nama_bulan}}' in run.text:
                found_month_placeholder = True
                new_text = run.text.replace('{{nama_bulan}}', month)
                run.text = new_text
                run.font.bold = False
                run.font.italic = False

                if is_first_page:
                    run.text = new_text.upper()  # Set teks nama bulan menjadi huruf kapital
                    run.font.size = Pt(16)  # Set ukuran font menjadi 16 di halaman pertama
                    print("Nama bulan di halaman pertama diganti dengan font size 16")
                else:
                    run.text = new_text  # Tidak ada perubahan kapitalisasi di halaman selanjutnya
                    run.font.size = Pt(12)  # Set ukuran font menjadi 12 di halaman setelahnya
                    print("Nama bulan di halaman selanjutnya diganti dengan font size 12")
                run.font.name = 'Arial'

        element_count += 1
        if element_count > 5:  # Asumsi jika lebih dari 5 elemen, pindah ke halaman kedua
            is_first_page = False

    # Mengganti bulan di tabel
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if '{{nama_bulan}}' in run.text:
                            found_month_placeholder = True
                            new_text = run.text.replace('{{nama_bulan}}', month)
                            run.text = new_text
                            run.font.bold = False
                            run.font.italic = False

                            if is_first_page:
                                run.text = new_text.upper()  # Set teks nama bulan menjadi huruf kapital
                                run.font.size = Pt(16)  # Set ukuran font menjadi 16 di halaman pertama
                                print("Nama bulan di halaman pertama (tabel) diganti dengan font size 16")
                            else:
                                run.text = new_text  # Tidak ada perubahan kapitalisasi di halaman selanjutnya
                                run.font.size = Pt(12)  # Set ukuran font menjadi 12 di halaman setelahnya
                                print("Nama bulan di halaman selanjutnya (tabel) diganti dengan font size 12")
                            run.font.name = 'Arial'
            element_count += 1
            if element_count > 5:  # Asumsi jika lebih dari 5 elemen, pindah ke halaman kedua
                is_first_page = False

    if not found_month_placeholder:
        print(f"Placeholder '{{nama_bulan}}' tidak ditemukan dalam dokumen.")  # Debugging



# Fungsi untuk memproses setiap folder dan mengisi tabel yang ada
def insert_dates_and_places_in_existing_table(doc_path, main_folder_name):
    doc = Document(doc_path)
    all_dates = []
    all_places = []
    month_set = set()

    folder_list = [folder_name for folder_name in os.listdir(main_folder_name) if os.path.isdir(os.path.join(main_folder_name, folder_name))]

    print("Daftar folder yang ditemukan:")
    for folder in folder_list:
        print(f"- {folder}")

    for folder_name in folder_list:
        print(f"Memproses folder: {folder_name}")
        current_date, current_place, month = get_date_and_place_from_folder(folder_name)

        if current_date and current_place:
            all_dates.append(current_date)
            all_places.append(current_place)
            month_set.add(month)
        else:
            print(f"Folder '{folder_name}' tidak terbaca dengan benar.")

    print(f"Total tanggal yang ditemukan: {len(all_dates)}")

    fill_existing_table(doc, all_dates, all_places)
    replace_placeholder_in_paragraphs(doc, len(folder_list))

    output_month = next(iter(month_set)) if month_set else "output"
    new_doc_path = f'Cover {output_month}.docx'
    
    # Ganti nama bulan di paragraf
    replace_month_names(doc, output_month)

    doc.save(new_doc_path)
    
    print(f"Dokumen berhasil disimpan sebagai {new_doc_path}")

# Main function
if __name__ == "__main__":
    monthFolder = input("Input nama folder untuk dijadikan cover: ")
    doc_path = 'COVER .docx'  # Asumsikan file .docx ada di direktori yang sama
    insert_dates_and_places_in_existing_table(doc_path, monthFolder)